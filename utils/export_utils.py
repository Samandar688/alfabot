import csv
import io
from typing import List, Dict, Any
from datetime import datetime
import asyncio
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportUtils:
    """Utility class for exporting data in various formats"""
    
    @staticmethod
    def generate_csv(data: List[Dict[str, Any]], filename: str = None) -> io.StringIO:
        """Generate CSV format from data"""
        if not data:
            return io.StringIO()
            
        output = io.StringIO()
        if data:
            fieldnames = data[0].keys()
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)
        
        output.seek(0)
        return output
    
    @staticmethod
    def generate_excel(data: List[Dict[str, Any]], sheet_name: str = "Data", title: str = None) -> io.BytesIO:
        """Generate Excel format from data"""
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        current_row = 1
        
        # Add title if provided
        if title:
            ws.cell(row=current_row, column=1, value=title)
            title_cell = ws.cell(row=current_row, column=1)
            title_cell.font = Font(size=16, bold=True)
            title_cell.alignment = Alignment(horizontal='center')
            title_cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            current_row += 2  # Skip a row after title
        
        if data:
            # Add headers
            headers = list(data[0].keys())
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=col_num, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
            
            current_row += 1
            
            # Add data rows
            for row_data in data:
                for col_num, value in enumerate(row_data.values(), 1):
                    ws.cell(row=current_row, column=col_num, value=value)
                current_row += 1
            
            # Auto-adjust column widths
            for col_num in range(1, len(headers) + 1):
                max_length = 0
                
                # Check header length
                header_length = len(str(headers[col_num - 1]))
                if header_length > max_length:
                    max_length = header_length
                
                # Check data lengths in this column
                for row_num in range(1, current_row):
                    try:
                        cell_value = ws.cell(row=row_num, column=col_num).value
                        if cell_value is not None:
                            cell_length = len(str(cell_value))
                            if cell_length > max_length:
                                max_length = cell_length
                    except:
                        continue
                
                # Set column width
                adjusted_width = min(max_length + 2, 50)
                column_letter = chr(64 + col_num)  # Convert to letter (A, B, C, etc.)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output
    
    @staticmethod
    def generate_word(data: List[Dict[str, Any]], title: str = "Export Hisoboti") -> io.BytesIO:
        """Generate Word document from data"""
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_paragraph = doc.add_paragraph(f"Yaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        if data:
            # Create table
            headers = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header).replace('_', ' ').title()
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Add data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data.values()):
                    row_cells[i].text = str(value) if value is not None else ""
        else:
            doc.add_paragraph("Export uchun ma'lumotlar mavjud emas.")
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    @staticmethod
    def generate_pdf(data: List[Dict[str, Any]], title: str = "Export Hisoboti") -> io.BytesIO:
        """Generate PDF document from data"""
        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=A4)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        story = []
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Yaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        if data:
            # Prepare table data
            headers = list(data[0].keys())
            table_data = [headers]
            
            for row in data:
                table_data.append([str(value) if value is not None else "" for value in row.values()])
            
            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
        else:
            story.append(Paragraph("Export uchun ma'lumotlar mavjud emas.", styles['Normal']))
        
        doc.build(story)
        output.seek(0)
        return output
    
    @staticmethod
    def get_filename_with_timestamp(base_name: str, extension: str) -> str:
        """Generate filename with timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base_name}_{timestamp}.{extension}"
    
    @staticmethod
    def format_data_for_export(data: List[Dict[str, Any]], export_type: str = "inventory") -> List[Dict[str, Any]]:
        """Format data specifically for export based on type"""
        if not data:
            return []
        
        formatted_data = []
        
        for item in data:
            if export_type == "inventory":
                formatted_item = {
                    "ID": item.get('id', ''),
                    "Nomi": item.get('name', ''),
                    "Miqdori": item.get('quantity', 0),
                    "Narxi": f"{item.get('price', 0):.2f}" if item.get('price') else "0.00",
                    "Umumiy Qiymat": f"{(item.get('quantity', 0) * item.get('price', 0)):.2f}" if item.get('price') else "0.00",
                    "Seriya Raqami": item.get('serial_number', ''),
                    "Tavsifi": item.get('description', ''),
                    "Yaratilgan": item.get('created_at', '').strftime('%Y-%m-%d %H:%M:%S') if item.get('created_at') else '',
                    "Yangilangan": item.get('updated_at', '').strftime('%Y-%m-%d %H:%M:%S') if item.get('updated_at') else ''
                }
            elif export_type == "statistics":
                formatted_item = {
                    "Ko'rsatkich": item.get('metric', ''),
                    "Qiymat": item.get('value', ''),
                    "Davr": item.get('period', ''),
                    "Sana": item.get('date', '').strftime('%Y-%m-%d') if item.get('date') else ''
                }
            else:
                formatted_item = item
            
            formatted_data.append(formatted_item)
        
        return formatted_data